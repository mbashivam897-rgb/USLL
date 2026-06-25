"""Build a professional Word (.docx) equity research report on United Spirits Ltd.
Mirrors United_Spirits_Equity_Research_Report.md. Charts in report/charts/.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
CH = os.path.join(HERE, "charts")

NAVY = RGBColor(0x0B, 0x25, 0x45)
BLUE = RGBColor(0x13, 0x55, 0x8C)
GOLD = RGBColor(0x8A, 0x6D, 0x0B)
GREY = RGBColor(0x55, 0x5B, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0xB3, 0x32, 0x2C)
GREEN = RGBColor(0x2E, 0x7D, 0x52)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.08

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=9, align="left", white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = WHITE
    elif color is not None:
        run.font.color.rgb = color

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    if level == 1:
        run.font.size = Pt(15); run.font.color.rgb = NAVY
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "3"); bottom.set(qn("w:color"), "0B2545")
        pbdr.append(bottom); pPr.append(pbdr)
    else:
        run.font.size = Pt(11.5); run.font.color.rgb = BLUE
    return p

def add_para(text, size=10, bold=False, italic=False, color=None, align="left", space_after=6):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT, "just": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    if color is not None: run.font.color.rgb = color
    return p

def add_bullet(runs, num=False):
    """runs: list of (text, bold) tuples"""
    p = doc.add_paragraph(style="List Number" if num else "List Bullet")
    p.paragraph_format.space_after = Pt(3)
    for text, bold in runs:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(10)
    return p

def add_table(headers, rows, widths=None, header_bg="0B2545", first_col_bold=True, font=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], header_bg)
        set_cell_text(hdr[i], h, bold=True, white=True, size=font,
                      align="left" if i == 0 else "center")
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        if r_idx % 2 == 1:
            for c in cells: set_cell_bg(c, "EEF2F7")
        for i, val in enumerate(row):
            bold = first_col_bold and i == 0
            set_cell_text(cells[i], val, bold=bold, size=font,
                          align="left" if i == 0 else "center")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def add_chart(fname, caption, width=6.3):
    path = os.path.join(CH, fname)
    if not os.path.exists(path): return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    run = c.add_run(caption); run.italic = True; run.font.size = Pt(8); run.font.color.rgb = GREY

# ============================================================= COVER PAGE
sec = doc.sections[0]
sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)

for _ in range(2):
    doc.add_paragraph()
add_para("EQUITY RESEARCH  |  INITIATING COVERAGE", size=11, bold=True, color=GOLD, align="center", space_after=2)
add_para("Consumer Staples, Beverages (Spirits)  •  India", size=10, color=GREY, align="center", space_after=14)
add_para("UNITED SPIRITS LIMITED", size=30, bold=True, color=NAVY, align="center", space_after=2)
add_para("a Diageo plc Group company", size=12, italic=True, color=GREY, align="center", space_after=16)

# recommendation banner
banner = doc.add_table(rows=1, cols=3); banner.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (lbl, val, clr) in enumerate([
    ("RECOMMENDATION", "HOLD / NEUTRAL", NAVY),
    ("TARGET PRICE", "₹1,348", GREEN),
    ("UPSIDE", "+1.7%", GREEN)]):
    cell = banner.rows[0].cells[i]
    set_cell_bg(cell, "0B2545" if i == 0 else "EEF2F7")
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = cell.paragraphs[0].add_run(lbl + "\n"); r1.font.size = Pt(8); r1.bold = True
    r1.font.color.rgb = WHITE if i == 0 else GREY
    r2 = cell.paragraphs[0].add_run(val); r2.font.size = Pt(14); r2.bold = True
    r2.font.color.rgb = WHITE if i == 0 else clr
    cell.width = Inches(2.2)
doc.add_paragraph().paragraph_format.space_after = Pt(8)

# snapshot table
snap = [
    ("Tickers", "NSE: UNITDSPR  •  BSE: 532432"),
    ("Current Market Price (CMP)", "₹1,326"),
    ("12-Month Target Price", "₹1,348"),
    ("52-Week Range", "₹1,210 - ₹1,489"),
    ("Market Capitalisation", "₹96,453 crore (~US$11.3 bn)"),
    ("Enterprise Value", "₹93,726 crore"),
    ("Shares Outstanding", "72.74 crore (Face value ₹2)"),
    ("Free Float", "~44% (Diageo holds ~56%)"),
    ("Valuation Basis", "60% EV/EBITDA, 20% P/E, 5% EV/EBIT, 15% DCF"),
    ("Investment Horizon", "12 months"),
]
t = doc.add_table(rows=0, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
for i, (k, v) in enumerate(snap):
    cells = t.add_row().cells
    set_cell_bg(cells[0], "DCE3EC")
    set_cell_text(cells[0], k, bold=True, size=9.5, color=NAVY)
    set_cell_text(cells[1], v, size=9.5)
    cells[0].width = Inches(2.6); cells[1].width = Inches(4.0)
doc.add_paragraph().paragraph_format.space_after = Pt(6)
add_para("Analyst: Senior Equity Research Analyst, CFA", size=10, bold=True, align="center", space_after=1)
add_para("Date of Publication: 25 June 2026", size=10, color=GREY, align="center", space_after=10)

# disclaimer box
d = doc.add_table(rows=1, cols=1); d.style = "Table Grid"; d.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_bg(d.rows[0].cells[0], "FBF6E9")
dc = d.rows[0].cells[0]
dp = dc.paragraphs[0]; dp.paragraph_format.space_after = Pt(2)
rr = dp.add_run("Disclaimer:  "); rr.bold = True; rr.font.size = Pt(8); rr.font.color.rgb = GOLD
rr2 = dp.add_run("Prepared for academic / capstone and illustrative research purposes per CFA Institute research-reporting standards. "
                 "This is NOT investment advice and NOT a solicitation to buy or sell any security. Financial forecasts and valuation outputs "
                 "are sourced from the uploaded model (United Spirits Excell.xlsx) and the United Spirits FY2025 Integrated Annual Report; "
                 "market/industry/macro context is from public-domain references cited inline. Third-party content has been paraphrased for "
                 "compliance with licensing restrictions.")
rr2.font.size = Pt(8); rr2.font.color.rgb = GREY

doc.add_page_break()

# ============================================================= 1. INVESTMENT SUMMARY
add_heading("1.  Investment Summary", 1)
add_para("We initiate coverage on United Spirits Limited (USL) with a HOLD / NEUTRAL rating and a 12-month target price of ₹1,348, "
         "implying ~1.7% upside to the current market price of ₹1,326. USL is, in our view, the highest-quality large-cap proxy for "
         "India's structural spirits-premiumisation theme, but after a multi-year re-rating, the current valuation already discounts the bulk "
         "of that optimism. The risk-reward is balanced; we would accumulate on weakness below ~₹1,150 and trim above ~₹1,500.", align="just")

add_heading("Investment Thesis (Key Points)", 2)
add_bullet([("Premiumisation is the durable growth engine. ", True),
            ("Prestige & Above (P&A) brands now contribute ~89-90% of net sales value (NSV), up from ~70% five years ago. With Indian premium/luxury "
             "spirits growing at low-double-digit rates vs mid-single-digit mainstream, USL can sustain ~9-10% NSV growth even with only mid-single-digit volume growth.", False)])
add_bullet([("Unrivalled brand portfolio and Diageo backing. ", True),
            ("A full-ladder portfolio, McDowell's No.1, Royal Challenge, Signature, Antiquity through Black Dog, Black & White, Godawan, Johnnie Walker, "
             "Singleton, Tanqueray, Don Julio, Smirnoff, backed by parent Diageo plc (~56%).", False)])
add_bullet([("Structural margin expansion. ", True),
            ("EBITDA margin rose from ~13% (FY23) to ~18% (FY26) and is modelled to reach ~20% by FY30E on mix, operating leverage and productivity.", False)])
add_bullet([("Fortress balance sheet and high returns. ", True),
            ("Net cash (net debt/EBITDA ~ -1.2x), ROCE/ROE ~20-23%, asset-light model, ~70% dividend payout.", False)])
add_bullet([("Crystallised optionality, the RCB monetisation. ", True),
            ("March 2026: USL completed the ₹16,660 crore all-cash sale of the Royal Challengers Bengaluru franchise to an Aditya Birla-led consortium, "
             "value not captured in our model-derived target and a source of potential capital return / re-rating.", False)])

add_heading("Why HOLD rather than BUY", 2)
add_para("Our blended fair value of ₹1,348 sits essentially level with the market price. The DCF value of ₹642 is well below the market price; the "
         "valuation is sustained almost entirely by rich relative multiples (~37-39x forward EV/EBITDA, ~52x trailing P/E). This is justified by quality "
         "and growth but leaves a thin margin of safety, the business case and the valuation case diverge.", align="just")

add_heading("Valuation Summary (per uploaded model)", 2)
add_table(["Methodology", "Implied Value/Share (₹)", "Weight", "Contribution (₹)"],
          [["EV/EBITDA (FY27E, peer median)", "1,411", "60%", "846"],
           ["P/E (FY27E, peer median)", "1,675", "20%", "335"],
           ["EV/EBIT (FY27E, peer median)", "1,408", "5%", "70"],
           ["DCF (FCFF, 2-stage)", "642", "15%", "96"],
           ["Blended Target Price", "", "100%", "1,348"]],
          widths=[2.9, 1.7, 0.9, 1.3])
add_chart("ex6_football_field.png", "Exhibit 6: Valuation summary, implied value per methodology & blended target.")

add_heading("Key Risks", 2)
add_para("Adverse state excise/tax actions; pricing-approval lags vs input inflation; intensifying premium competition (Pernod Ricard, Radico Khaitan, a "
         "recapitalised Tilaknagar Industries, Allied Blenders); customer/geographic concentration (a single state corporation ~30% of revenue); ENA/grain, "
         "glass and scotch-concentrate cost volatility; valuation de-rating risk; and ESG/health-policy headwinds.", align="just")

doc.add_page_break()

# ============================================================= 2. COMPANY OVERVIEW
add_heading("2.  Company Overview", 1)
add_heading("2.1  History and Ownership", 2)
add_para("United Spirits Limited is India's largest beverage-alcohol (alcobev) company and the local operating company of Diageo plc, the world's leading "
         "premium-spirits group. USL traces its lineage to the McDowell & Co. and UB spirits businesses consolidated under the erstwhile UB Group; Diageo "
         "acquired control in 2013 (a landmark ~US$1.9 bn transaction, then India's largest in the sector) and currently holds ~56% of equity, with the balance "
         "~44% free float. The company is listed on the NSE (UNITDSPR) and BSE (532432), with 72.74 crore shares (₹2 face value).", align="just")

add_heading("2.2  Business Segments", 2)
add_bullet([("Beverage Alcohol (~96%+ of net sales): ", True), ("the core IMFL franchise, whisky, brandy, rum, vodka and gin across the full price ladder.", False)])
add_bullet([("Sports (discontinued): ", True), ("the RCB IPL/WPL franchise via RCSPL, classified held-for-sale and fully divested in March 2026 (see §2.7).", False)])
add_para("Within beverage alcohol: Prestige & Above (P&A), the growth/margin driver, was ~89% of NSV in FY25; Popular is mass-market and deliberately de-emphasised "
         "(a tranche was franchised/slump-sold in FY2023).", align="just")

add_heading("2.3  Product Portfolio, Key Brands", 2)
add_table(["Tier", "Representative Brands"],
          [["Lower / Mid Prestige", "McDowell's No.1 (whisky/brandy/rum), Royal Challenge"],
           ["Upper Prestige", "Signature, Antiquity, Royal Challenge American Pride"],
           ["Premium Scotch", "Black & White (fastest-growing primary Scotch globally; India largest market), Black Dog"],
           ["Luxury / Imported (Diageo)", "Johnnie Walker, The Singleton, Don Julio (tequila), Tanqueray (gin), Cîroc, Ketel One, Captain Morgan, Baileys"],
           ["Indian Craft / Luxury Malt", "Godawan (India's most-awarded artisanal single malt, 85+ accolades)"],
           ["White Spirits / New-to-World", "Smirnoff (incl. flavours), McDowell's X Series (gin/vodka/rum)"],
           ["Craft (via Ventures)", "Greater Than & Hāpusa gin (NAO Spirits)"]],
          widths=[2.0, 4.6])

add_heading("2.4  Manufacturing Footprint", 2)
add_para("Asset-light and geographically dispersed: owned distilleries/bottling units plus tie-up/contract manufacturing and franchisee bottlers across ~40+ facilities. "
         "Because alcohol is a State subject with inter-state movement restrictions, production is largely in-state. Asset turnover (~7x) and capex (~1-2% of NSV) reflect "
         "this capital-light structure; an ongoing supply-agility programme drives structural cost savings.", align="just")

add_heading("2.5  Distribution Network", 2)
add_para("USL reaches 70,000+ retail outlets. India's route-to-market is state-controlled: in many states USL sells to government beverage corporations (e.g., TASMAC), "
         "creating customer concentration, a single state customer is ~30% of revenue, a structural risk typical of the industry.", align="just")

add_heading("2.6  The Diageo Relationship", 2)
add_para("As Diageo's India arm, USL gains access to the global luxury portfolio, R&D/marketing and responsible-drinking frameworks, conservative treasury/governance, and a "
         "global premiumisation playbook. India is a structural bright spot within Diageo's portfolio at a time the parent has cut global guidance (see §5 and §12).", align="just")

add_heading("2.7  Recent Strategic Developments", 2)
add_bullet([("RCB divestiture (completed March 2026): ", True),
            ("USL sold 100% of RCSPL for ₹16,660 crore (~US$1.7-1.9 bn) all-cash to a consortium led by the Aditya Birla Group (with The Times of India Group, Bolt Ventures, "
             "Blackstone/BXPE). The franchise was ~17% of market cap; brokerages widely viewed the exit as value-unlocking.", False)])
add_bullet([("NAO Spirits acquisition / Ventures: ", True),
            ("Through Diageo India Ventures, USL invested in craft start-ups, NAO Spirits (Greater Than, Hāpusa gin), Pistola, Sober, Quaffine (FY25: ₹18 cr secured loan to NAO "
             "plus equity). Rationale: low-cost early exposure to fast-growing craft/new-to-world categories, strengthening the premiumisation funnel toward younger consumers.", False)])

doc.add_page_break()

# ============================================================= 3. INDUSTRY
add_heading("3.  Industry Analysis", 1)
add_heading("3.1  Market Structure and Size", 2)
add_para("India is the world's third-largest alcobev market by volume and the largest whisky market. Total IMFL volumes are ~410 million cases (2024, IWSR) and the broader "
         "alcohol market roughly US$50-70 billion, growing ~6% CAGR. The IMFL spirits market is an oligopoly: USL (Diageo) and Pernod Ricard India lead, followed by Allied "
         "Blenders, Radico Khaitan, Tilaknagar Industries and Globus Spirits, plus regional players.", align="just")

add_heading("3.2  Porter's Five Forces", 2)
add_table(["Force", "Intensity", "Assessment"],
          [["Threat of New Entrants", "Low", "High barriers: 30+ state excise regimes, capital/working-capital intensity (maturation), advertising bans, entrenched distribution."],
           ["Supplier Power", "Moderate", "ENA/grain, scotch concentrate, glass and packaging can spike (glass/ENA hit FY23). Scale + Diageo sourcing offset."],
           ["Buyer Power", "High", "State corporations control pricing/RTM and cap MRP; ~30% revenue concentration in one customer amplifies this."],
           ["Threat of Substitutes", "Moderate", "Beer, wine, country liquor, low/no-alcohol; premiumisation shifts consumers toward branded premium spirits."],
           ["Competitive Rivalry", "High & rising", "Pernod, Radico, ABDL and a recapitalised Tilaknagar (Imperial Blue). USL competes on value share and brand equity."]],
          widths=[1.7, 1.2, 3.7])
add_para("Net assessment: high entry barriers and structural premiumisation are attractive, but buyer power (state control) and rising rivalry cap pricing freedom and returns, "
         "supporting premium-but-not-unlimited multiples.", italic=True, align="just")

add_heading("3.3  Premiumisation, The Structural Growth Driver", 2)
add_para("P&A IMFL and premium beer are growing at low-double-digit rates vs mid-single-digit mainstream (JM Financial, Sep-2025). Luxury spirits are only ~3% of cases sold but "
         ">20% of industry profits (IWSR). Financial implication for USL: with a ~89-90% P&A mix, each point of mix enrichment and price/mix gain flows disproportionately to gross "
         "and EBITDA margins, the mechanism behind our modelled ~46%→48% gross and ~18%→20% EBITDA margin path. Value growth, not volume growth, is the engine.", align="just")
add_chart("ex3_volume_mix.png", "Exhibit 3: Volume mix shift, deliberate exit of 'Popular'; P&A now the overwhelming majority of volume and value.", width=6.0)

add_heading("3.4  State Regulation, Taxation and Distribution Barriers", 2)
add_para("Alcohol sits outside GST and is a State subject: each state controls licensing, pricing approvals, label registration, RTM and excise duty (a pass-through). Hence the topline is "
         "best tracked on Net Sales Value (NSV = revenue net of excise). Key 2025-26 developments:", align="just")
add_bullet([("Karnataka Alcohol-in-Beverage (AIB) excise (effective 11 May 2026): ", True),
            ("first Indian state to adopt alcohol-content-based duty within rationalised slabs, deregulating price-fixation so producers set prices by market/alcohol content; targets ~₹45,000 cr "
             "excise revenue, transitioning over 3-4 years. As Karnataka is ~17% of national whisky volumes, greater pricing freedom is a structural positive for premium players, though the "
             "transition adds near-term uncertainty.", False)])
add_bullet([("Prohibition / RTM risk: ", True), ("dry states (Gujarat, Bihar) and periodic RTM reversals (Delhi, Andhra Pradesh) remain the principal regulatory tail risks.", False)])

doc.add_page_break()

# ============================================================= 4. MACRO
add_heading("4.  Economic and Macro Analysis", 1)
add_table(["Macro Indicator", "Reading", "Source"],
          [["India real GDP growth FY26", "+7.7% (Q4 +7.8%); nominal +8.9%", "MoSPI / Economic Survey"],
           ["FY27 GDP projection", "6.8-7.2% (Economic Survey); RBI ~7.4%", "PRS / RBI"],
           ["Status", "Fastest-growing major economy (4th yr)", "Economic Survey 2025-26"],
           ["Demographics", "Large young legal-drinking-age base; rising urbanisation", "n/a"]],
          widths=[2.3, 3.0, 1.3])
add_para("Transmission to spirits demand. India's combination of ~9% nominal GDP growth, an expanding middle class, rising disposable incomes, rapid urbanisation, and a young aspirational "
         "legal-drinking-age cohort creates a powerful structural tailwind for value-led alcobev consumption. Per-capita consumption is low by global standards, leaving a long runway. Growth is "
         "increasingly 'drinking better, not more', favouring USL's premium mix. The main macro risk is a discretionary-spending slowdown (high food inflation or weak urban sentiment) prompting "
         "down-trading that disproportionately affects the upper-premium tier.", align="just")

# ============================================================= 5. COMPETITIVE
add_heading("5.  Competitive Positioning", 1)
add_para("USL is the #1 spirits company in India by volume and a value leader. Below we benchmark USL against listed peers on the metrics in the uploaded model (FY26 actuals; United Breweries as a "
         "beer reference). Tilaknagar Industries is discussed qualitatively (not in the model's comparable set).", align="just")
add_heading("5.1  Peer Comparison (FY2026)", 2)
add_table(["Company", "NSV (₹cr)", "EBITDA (₹cr)", "EBITDA %", "Mkt Cap (₹cr)", "EV (₹cr)", "EV/EBITDA", "P/E"],
          [["United Spirits", "12,467", "2,279", "18.3%", "96,453", "93,726", "41.1", "52.5"],
           ["Radico Khaitan", "6,050", "1,038", "17.1%", "50,274", "50,883", "49.0", "81.5"],
           ["United Breweries (beer ref.)", "9,240", "824", "8.9%", "35,605", "36,350", "44.1", "86.2"],
           ["Allied Blenders & Distillers", "3,923", "542", "13.8%", "17,640", "18,635", "34.4", "77.0"],
           ["Peer median", "n/a", "n/a", "n/a", "n/a", "n/a", "44.1", "81.5"]],
          widths=[1.9, 0.8, 0.85, 0.7, 0.95, 0.85, 0.9, 0.6], font=8)
add_para("ROE/ROCE: USL ~20-23% (model), the highest-quality return profile in the peer set on a scale-adjusted basis.", italic=True, size=9)
add_chart("ex5_peer_multiples.png", "Exhibit 5: Peer valuation (FY26), USL trades below premium peers on P/E.", width=6.0)

add_heading("5.2  Competitor Assessment", 2)
add_bullet([("Radico Khaitan, ", True), ("fastest-premiumising challenger; strong FY26 (Q3 EBITDA margin ~17%, gross ~47%; luxury value growth ~50% YoY) on Rampur, Jaisalmer, After Dark. Trades at a premium (~49x EV/EBITDA).", False)])
add_bullet([("Allied Blenders (ABDL), ", True), ("record FY26 (revenue +11.5%, EBITDA +28.5% to ₹568 cr, margin +163 bps to 14.4%; adj. PAT +36%), pushing premiumisation (ICONiC). Lower-margin but improving fast.", False)])
add_bullet([("Tilaknagar Industries, ", True), ("brandy leader (Mansion House) that acquired Pernod's Imperial Blue whisky for ~₹4,150 cr EV (completed Dec-2025), India's largest alcobev deal since Diageo-USL (2013), intensifying mid-premium whisky competition.", False)])
add_bullet([("United Breweries (reference), ", True), ("India's beer leader (Heineken-controlled); a substitute/wallet competitor rather than a direct spirits peer.", False)])
add_bullet([("Diageo plc (parent), ", True), ("supplies USL's premium/luxury portfolio and global capabilities; its global softness contrasts with India's strength.", False)])
add_para("Positioning verdict: USL has the broadest portfolio, deepest distribution, strongest returns (ROCE ~22%) and most defensive (net-cash) balance sheet, but cedes some top-line growth optics to Radico and "
         "faces a re-energised Tilaknagar. Its value-share-over-volume-share strategy is the right one for a premiumising market.", italic=True, align="just")

doc.add_page_break()

# ============================================================= 6. FINANCIAL ANALYSIS
add_heading("6.  Financial Analysis (Historical, FY2017-FY2026)", 1)
add_para("All figures consolidated, NSV (net-of-excise) basis, ₹ crore. Source: United Spirits Excell.xlsx and FY2025 Integrated Annual Report.", italic=True, size=9)
add_heading("6.1  Income Statement Trends", 2)
add_table(["Metric", "FY22", "FY23", "FY24", "FY25", "FY26"],
          [["Net Sales Value (NSV)", "9,712", "10,612", "11,321", "12,069", "12,467"],
           ["NSV growth %", "n/a", "+9.3%", "+6.7%", "+6.6%", "+3.3%"],
           ["Gross profit", "4,032", "4,109", "4,779", "5,209", "5,793"],
           ["Gross margin %", "42%", "39%", "42%", "43%", "46%"],
           ["EBITDA", "1,595", "1,417", "2,000", "2,236", "2,279"],
           ["EBITDA margin %", "16%", "13%", "18%", "19%", "18%"],
           ["EBIT", "1,291", "1,134", "1,725", "1,953", "1,990"],
           ["PAT (reported)", "811", "1,126", "1,408", "1,582", "1,838"],
           ["PAT margin %", "8%", "11%", "12%", "13%", "15%"]],
          widths=[2.2, 0.9, 0.9, 0.9, 0.9, 0.9])
add_para("Read-through. NSV grew every year while total volume fell (from ~79 mn cases in FY22 to ~64 mn in FY25), a direct consequence of the Popular exit and premiumisation; topline is now realisation-led. "
         "Gross margin recovered sharply from the FY23 input-cost trough (39%) to ~46% by FY26. FY26 results were corroborated by reported numbers, consolidated NSV ₹12,467 cr (+7.7%), Q4 PAT +28% YoY, Q4 EBITDA beating estimates.", align="just")
add_chart("ex1_nsv_growth.png", "Exhibit 1: Net Sales Value & growth, premiumisation-led topline (FY22-FY30E).")
add_chart("ex2_margins.png", "Exhibit 2: Margin trajectory, structural expansion on mix & operating leverage.")

add_heading("6.2  Balance Sheet, Leverage and Capital Structure", 2)
add_table(["Metric", "FY22", "FY23", "FY24", "FY25", "FY26"],
          [["Debt/Equity (x)", "0.13", "0.04", "0.05", "0.07", "0.05"],
           ["Net debt/EBITDA (x)", "-0.01", "-0.67", "-0.81", "-1.08", "-1.20"],
           ["Current ratio (x)", "1.42", "1.76", "1.93", "2.00", "2.07"],
           ["Interest coverage (x)", "14.7", "10.9", "22.7", "21.9", "12.6"],
           ["Cash & treasury (₹cr)", "282", "1,139", "1,868", "2,903", "3,134"]],
          widths=[2.2, 0.9, 0.9, 0.9, 0.9, 0.9])
add_para("USL deleveraged to a net-cash position (~₹3,134 cr at FY26, net debt/EBITDA ~ -1.2x), reflecting strong cash generation and Diageo's conservative treasury. This fortress balance sheet underpins the dividend "
         "and provides ample flexibility, now augmented by ₹16,660 cr of RCB proceeds.", align="just")

add_heading("6.3  Working Capital", 2)
add_para("Steady-state assumptions: receivables ~97 days, inventory ~135 days, payables ~104 days. Inventory is structurally high (whisky maturation); receivables are elevated by dependence on state corporations. This is a "
         "working-capital-intensive model and the principal drag on free-cash conversion.", align="just")
add_heading("6.4  Cash Flow and Cash Conversion", 2)
add_para("Operating cash flow was ~₹1,459 cr in FY26; capex is consistently small (~₹180-330 cr p.a.). Healthy CFO, low capex and ~70% dividend payout make USL a reliable cash compounder, though working-capital build in growth "
         "years moderates free-cash conversion in any given year.", align="just")
add_chart("ex9_fcf_div.png", "Exhibit 9: Free cash flow vs dividends (~70% payout).", width=5.6)

doc.add_page_break()

# ============================================================= 7. FORECAST
add_heading("7.  Forecast Analysis (FY2027E-FY2030E)", 1)
add_para("FY2026 is the last actual; forecasts are the continuing beverage-alcohol business per the uploaded model.", italic=True, size=9)
add_table(["₹ crore", "FY26", "FY27E", "FY28E", "FY29E", "FY30E"],
          [["Net Sales Value (NSV)", "12,467", "13,589", "14,812", "15,997", "17,117"],
           ["NSV growth %", "+3.3%", "+9.0%", "+9.0%", "+8.0%", "+7.0%"],
           ["EBITDA", "2,279", "2,547", "2,851", "3,159", "3,465"],
           ["EBITDA margin %", "18.3%", "18.7%", "19.2%", "19.7%", "20.2%"],
           ["EBIT", "1,990", "2,170", "2,433", "2,700", "2,965"],
           ["PAT", "1,838*", "1,709", "1,891", "2,119", "2,318"],
           ["EPS (₹)", "25.3*", "23.5", "26.0", "29.1", "31.9"]],
          widths=[2.2, 0.9, 0.9, 0.9, 0.9, 0.9])
add_para("*FY26 PAT/EPS include one-off/discontinued-operation items; FY27E reflects normalised continuing operations, hence the optical YoY dip.", italic=True, size=8.5)
add_heading("7.1  Revenue Drivers", 2)
add_bullet([("Premiumisation & mix: ", True), ("continued P&A salience gains (toward ~90%+ of NSV) and trade-up within categories drive realisation growth.", False)])
add_bullet([("Volume: ", True), ("mid-single-digit, now the Popular-exit drag has lapped.", False)])
add_bullet([("Pricing: ", True), ("state-by-state increases; Karnataka AIB deregulation could improve pricing freedom over time.", False)])
add_bullet([("New products: ", True), ("white spirits (Smirnoff/X Series), luxury malts (Godawan), craft (NAO) extend the funnel.", False)])
add_heading("7.2  Margin Drivers", 2)
add_para("EBITDA margin expands ~40-50 bps p.a. (18.3%→~20.2%) on mix enrichment (gross margin ~46%→48%), operating leverage, and supply-agility savings (partly reinvested in A&P). We stop short of the >22% margins implied by bulls, "
         "given state-gated pricing and rising competition.", align="just")
add_heading("7.3  Key Model Assumptions", 2)
add_table(["Assumption", "Value", "Basis"],
          [["Effective tax rate", "25%", "Normalised"],
           ["Dividend payout", "70%", "Held at FY26 level"],
           ["Receivable / Inventory / Payable days", "97 / 135 / 104", "10-yr / 3-yr medians"],
           ["Capex", "~1-2% of NSV", "4-5 yr median"],
           ["Depreciation useful life", "16 years", "~1/6.2% avg"],
           ["Treasury yield", "market yield on deposits/MFs", "n/a"]],
          widths=[2.7, 1.6, 2.3])

doc.add_page_break()

# ============================================================= 8. VALUATION
add_heading("8.  Valuation", 1)
add_para("We adopt the uploaded model's blended methodology, emphasising relative multiples (forward-looking, market-anchored) with a DCF cross-check.", align="just")
add_heading("8.1  DCF Valuation (FCFF, 2-stage)", 2)
add_table(["Parameter", "Value"],
          [["Risk-free rate (10-yr G-Sec)", "7.0%"],
           ["Equity risk premium (Rm - Rf)", "4.0%"],
           ["Adjusted beta (0.67 × 0.89 + 0.33)", "0.92"],
           ["Cost of equity (CAPM) = WACC (debt negligible)", "11.0%"],
           ["Terminal growth (g)", "6.0%"],
           ["Sum of PV of explicit FCFF (FY27-34)", "₹11,269 cr"],
           ["PV of terminal value", "₹32,721 cr"],
           ["Enterprise Value", "₹43,990 cr"],
           ["(+) Cash / (-) Debt", "+₹3,134 cr / -₹407 cr"],
           ["Equity Value", "₹46,717 cr"],
           ["Shares outstanding", "72.74 cr"],
           ["DCF Value per Share", "₹642"]],
          widths=[4.2, 2.4])
add_para("Honest finding: even on supportive inputs, the DCF yields ~₹642, well below the market price. High working-capital intensity depresses free-cash conversion, and the market implicitly applies a lower discount rate / "
         "higher perpetual growth than our base case. The DCF therefore functions as a valuation floor, hence its 15% blend weight.", align="just")
add_chart("ex7_dcf_sensitivity.png", "Exhibit 7: DCF sensitivity, fair value/share (₹). Base case WACC 11% / g 6% → ₹642.", width=4.8)

add_heading("8.2  Relative Valuation", 2)
add_table(["Multiple", "Peer median (x)", "Implied Value/Share (₹)"],
          [["EV/EBITDA", "39.2", "1,411"],
           ["P/E", "71.3", "1,675"],
           ["EV/EBIT", "45.9", "1,408"],
           ["EV/Revenue (reference)", "4.4", "859"]],
          widths=[2.4, 2.0, 2.2])
add_para("Peer selection: Radico Khaitan, United Breweries and Allied Blenders. Premium/discount justification: USL trades at a discount to the fastest-growing premium peer (Radico) on P/E but a premium to mass-market peers, "
         "appropriate given USL's leadership, mix and superior returns, while its slower headline growth caps the premium. Multiples frame USL as fully, but not egregiously, valued.", align="just")

add_heading("8.3  Target Price Methodology and Weights", 2)
add_table(["Method", "Weight", "Rationale"],
          [["EV/EBITDA", "60%", "Capital-structure-neutral; the market's primary lens for staples; most reliable cross-peer metric."],
           ["P/E", "20%", "Captures earnings power / equity-investor lens."],
           ["EV/EBIT", "5%", "Secondary cross-check on operating value."],
           ["DCF", "15%", "Long-horizon intrinsic anchor; down-weighted due to terminal-value sensitivity and WC-driven FCF volatility."]],
          widths=[1.4, 0.9, 4.3])
add_para("Justification of the DCF/relative split (15% / 85%): (i) forecast visibility, multi-year cash flows for a state-regulated, WC-intensive business carry wide confidence intervals, making the DCF terminal value highly sensitive; "
         "(ii) industry maturity, a deep set of listed, actively-traded Indian alcobev comparables makes market multiples informative; (iii) reliability of market multiples, for a high-quality, widely-followed staple, traded multiples "
         "capture the premiumisation/quality premium better than a single discount-rate assumption. A more DCF-weighted analyst would arrive at a lower fair value.", align="just")
add_table(["Method", "Implied (₹)", "Weight", "Contribution (₹)"],
          [["EV/EBITDA", "1,411", "60%", "846"],
           ["P/E", "1,675", "20%", "335"],
           ["EV/EBIT", "1,408", "5%", "70"],
           ["DCF", "642", "15%", "96"],
           ["Blended Target Price", "", "", "1,348"]],
          widths=[2.0, 1.5, 1.0, 2.1])
add_heading("8.4  Upside not in the Target, RCB Proceeds", 2)
add_para("The model-derived ₹1,348 target does not capitalise the ₹16,660 crore RCB sale proceeds (~17% of market cap). Net of taxes/leakage and depending on redeployment (special dividend, buyback or core reinvestment), this is a "
         "material upside option and a key reason we rate HOLD rather than Reduce.", align="just")

doc.add_page_break()

# ============================================================= 9. ESG
add_heading("9.  ESG Analysis", 1)
add_para("USL's ESG agenda, 'Spirit of Progress' (aligned to Diageo's global framework), is a genuine competitive and risk-mitigation asset (FY2025 Integrated Annual Report).", align="just")
add_heading("9.1  Environmental", 2)
add_bullet([("Water: ", True), ("~54% improvement in water used per litre of spirit distilled vs FY2020; ~671k kilolitres consumed; 'Preserve Water for Life' stewardship, replenishment and regenerative agriculture (Godavari Initiative).", False)])
add_bullet([("Energy & emissions: ", True), ("99% renewable energy in direct operations (since FY2020); ~93% reduction in absolute Scope 1 & 2 GHG vs FY2020 (Scope 2 market-based ~zero after i-RECs).", False)])
add_bullet([("Net Zero roadmap: ", True), ("Net Zero in direct operations by 2030; across the value chain by 2050.", False)])
add_heading("9.2  Social", 2)
add_bullet([("Responsible drinking: ", True), ("~0.7 million people educated (Act Smart India on underage drinking; DRINKiQ).", False)])
add_bullet([("People: ", True), ("~2,400+ employees, ~86% retention, ~89% engagement, ~28% women's representation; active inclusion & diversity initiatives.", False)])
add_heading("9.3  Governance", 2)
add_bullet([("", False), ("Diageo governance standards and board oversight; robust corporate-governance framework; Diageo Marketing Code and consumer-information standards; conservative treasury policy.", False)])
add_heading("9.4  Impact on Valuation & Risk", 2)
add_para("Strong ESG execution reduces regulatory/reputational tail-risk in a sector under perennial health/policy scrutiny, supports premium-brand equity, and is consistent with the low beta in our WACC. We assign no explicit valuation "
         "premium but view USL's ESG leadership as supportive of multiple durability and a lower cost of capital over time.", align="just")

# ============================================================= 10. RISK
add_heading("10.  Risk Analysis", 1)
add_heading("10.1  Risk Matrix", 2)
add_table(["Risk", "Likelihood", "Impact", "Mitigant"],
          [["Regulatory / RTM disruption (policy reversals, prohibition)", "Medium", "High", "Diversification across 30+ states; Diageo compliance"],
           ["Excise / tax increases not passed through", "Med-High", "Med-High", "Premium mix, pricing power, AIB freedom (Karnataka)"],
           ["Raw-material inflation (ENA/grain, glass, scotch)", "Medium", "Medium", "Scale sourcing, supply-agility savings, productivity"],
           ["Competitive pressure (Pernod, Radico, Tilaknagar, ABDL)", "High", "Medium", "Portfolio breadth, distribution, value-share focus"],
           ["Consumer down-trading (macro slowdown)", "Medium", "Med-High", "Full price-ladder portfolio; mass-to-premium funnel"],
           ["Currency / import cost (scotch concentrate)", "Medium", "Low-Med", "Local production; partial hedging"],
           ["Customer concentration (~30% one state corp.)", "Structural", "Medium", "Industry-wide norm; receivables management"],
           ["ESG / health-policy headwinds", "Medium", "Medium", "ESG leadership; responsible-drinking programmes"],
           ["Valuation de-rating (rich multiple)", "Medium", "High", "Quality, growth, RCB cash optionality"]],
          widths=[2.6, 1.0, 0.9, 2.1], font=8)
add_heading("10.2  International / Global Risk Context", 2)
add_para("USL is a domestic IMFL business (limiting direct US-tariff exposure), but the global backdrop matters: IWSR data show global beverage-alcohol volumes fell ~2% in 2025 (third straight year), premiumisation has 'stalled'/fragmented "
         "amid macro pressure, and US tariffs cut Scotch exports to the US ~15% by volume. Parent Diageo cut FY26 guidance on US/China weakness and tariffs, reduced its dividend (Feb-2026), and is pursuing selective disposals (the RCB exit "
         "fits this). Implication: India stands out as a structural growth engine within a stagnant global spirits market, a relative positive for USL, though scotch-concentrate costs and parent-level strategy shifts bear watching.", align="just")

doc.add_page_break()

# ============================================================= 11. RECOMMENDATION
add_heading("11.  Investment Recommendation", 1)
add_table(["", ""],
          [["Recommendation", "HOLD / NEUTRAL"],
           ["12-Month Target Price", "₹1,348"],
           ["Current Market Price", "₹1,326"],
           ["Expected Total Return", "~+1.7% price + ~1.0% dividend yield ≈ ~3%"],
           ["Investment Horizon", "12 months"]],
          widths=[2.4, 4.2], first_col_bold=True)
add_para("Rationale. USL is a best-in-class, structurally advantaged compounder, the cleanest large-cap proxy for Indian spirits premiumisation, with a net-cash balance sheet, ~22% ROCE and a multi-year margin runway. However, the business "
         "case and the valuation case diverge: at ~37-39x forward EV/EBITDA and ~52x trailing earnings, the stock already discounts years of double-digit growth and margin expansion. Our blended fair value of ₹1,348 sits essentially at the "
         "market price, and the DCF floor (₹642) underscores how much rests on sustained premium multiples.", align="just")
add_para("Why USL deserves a premium to mass-market peers (but not to Radico). USL warrants a premium over ABDL/Globus for leadership, portfolio breadth, ~89-90% P&A mix, superior returns and balance-sheet strength. It trades at a discount to "
         "Radico Khaitan, whose faster top-line and luxury growth command a higher multiple, a gap we view as fair given USL's larger, more mature base.", align="just")
add_chart("ex8_scenarios.png", "Exhibit 8: Scenario-weighted target prices.", width=4.8)
add_table(["Scenario", "Key assumptions", "Target (₹)", "vs CMP"],
          [["Bull", "Vol +6%, realisation +6%, EBITDA margin →~24%, lower WACC, RCB cash redeployed accretively", "~1,760", "+33%"],
           ["Base", "Vol +4-4.5%, realisation +5%, margin →~20%, WACC 11%/g 6%", "1,348", "+1.7%"],
           ["Bear", "Vol +2.5%, realisation +4%, margin stalls ~17%, de-rating, adverse excise", "~920", "-31%"]],
          widths=[1.0, 3.8, 1.0, 0.8], font=8.5)
add_para("Action. Hold core positions for compounding and the un-modelled RCB-cash optionality. Accumulate below ~₹1,150; trim above ~₹1,500. Catalysts to monitor: redeployment of RCB proceeds (special dividend / buyback), Karnataka AIB "
         "pricing outcomes, P&A volume momentum, input-cost trajectory, and Diageo's August-2026 global strategy refresh.", align="just")

doc.add_page_break()

# ============================================================= 12. APPENDICES
add_heading("12.  Appendices", 1)
add_heading("Appendix A, Integrated Income Statement (₹ cr, NSV basis)", 2)
add_table(["", "FY22", "FY23", "FY24", "FY25", "FY26", "FY27E", "FY28E", "FY29E", "FY30E"],
          [["NSV", "9,712", "10,612", "11,321", "12,069", "12,467", "13,589", "14,812", "15,997", "17,117"],
           ["COGS", "5,681", "6,503", "6,542", "6,860", "6,674", "7,406", "7,999", "8,558", "9,072"],
           ["Gross profit", "4,032", "4,109", "4,779", "5,209", "5,793", "6,183", "6,814", "7,439", "8,045"],
           ["EBITDA", "1,595", "1,417", "2,000", "2,236", "2,279", "2,547", "2,851", "3,159", "3,465"],
           ["EBIT", "1,291", "1,134", "1,725", "1,953", "1,990", "2,170", "2,433", "2,700", "2,965"],
           ["PAT", "811", "1,126", "1,408", "1,582", "1,838", "1,709", "1,891", "2,119", "2,318"]],
          widths=[1.0]+[0.62]*9, font=7.5)
add_heading("Appendix B, Key Ratios", 2)
add_table(["", "FY24", "FY25", "FY26", "FY27E", "FY28E", "FY29E", "FY30E"],
          [["Gross margin", "42%", "43%", "46%", "46%", "46%", "47%", "47%"],
           ["EBITDA margin", "18%", "19%", "18%", "19%", "19%", "20%", "20%"],
           ["ROE", "21%", "21%", "22%", "19%", "19%", "20%", "21%"],
           ["ROCE", "23%", "23%", "21%", "22%", "22%", "23%", "23%"],
           ["Net debt/EBITDA", "-0.8", "-1.1", "-1.2", "-1.1", "-1.0", "-1.1", "-1.1"]],
          widths=[1.7]+[0.7]*7, font=8)
add_chart("ex4_returns.png", "Exhibit 4: Capital returns, sustained ~20%+ ROE / ROCE.")
add_heading("Appendix C, Cash Flow Summary (₹ cr)", 2)
add_table(["", "FY26", "FY27E", "FY28E", "FY29E", "FY30E"],
          [["CFO", "1,459", "1,292", "2,345", "1,957", "2,831"],
           ["Capex", "(181)", "(204)", "(222)", "(240)", "(257)"],
           ["Dividends paid", "(1,263)", "(1,145)", "(1,267)", "(1,420)", "(1,553)"]],
          widths=[1.8, 0.95, 0.95, 0.95, 0.95])
add_heading("Appendix D, Sources", 2)
add_para("Primary (per analyst mandate): United Spirits 'United Spirits Excell.xlsx' financial model (DCF, Comps, Valuation Summary, P&L, Balance Sheet, Cash Flow, Ratio Analysis, Beta); United Spirits Annual Reports FY2017, FY2019, FY2021, "
         "FY2023 and FY2025 Integrated Annual Report.", size=9, align="just")
add_para("Secondary (public domain, cited inline): Diageo India / Diageo plc investor communications; IWSR; The Hindu; Economic Times; Livemint/Mint; Financial Express; Hindustan Times; NDTV Profit; Reuters; Outlook Business; The Spirits "
         "Business; Quartr; PRS / Economic Survey 2025-26; MoSPI; RBI.", size=9, align="just")
add_para("All third-party content paraphrased/summarised for compliance with licensing restrictions. This document is an educational research report and does not constitute investment advice.", size=8.5, italic=True, color=GREY)

# ---- footer with page numbers ----
def add_page_number(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)

footer = doc.sections[0].footer
fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("United Spirits Limited (UNITDSPR)  |  Equity Research, HOLD, TP ₹1,348  |  Page ")
fr.font.size = Pt(7.5); fr.font.color.rgb = GREY
add_page_number(fp)

out = os.path.join(HERE, "United_Spirits_Equity_Research_Report.docx")
doc.save(out)
print("SAVED:", out, os.path.getsize(out), "bytes")
