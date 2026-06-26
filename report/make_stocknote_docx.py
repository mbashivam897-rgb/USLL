"""Build the United Spirits Ltd Stock Note (Word .docx), about 20 pages.
Style modelled on a sell-side Stock Note. Simple language. No em-dashes or tildes.
Data: United Spirits Excell.xlsx, fy 2026.xlsx, FY2025 Annual Report, public sources (cited).
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
CH = os.path.join(HERE, "charts")

NAVY = RGBColor(0x0B, 0x25, 0x45)
BLUE = RGBColor(0x13, 0x55, 0x8C)
GOLD = RGBColor(0x8A, 0x6D, 0x0B)
GREY = RGBColor(0x55, 0x5B, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0xB3, 0x32, 0x2C)
GREEN = RGBColor(0x2E, 0x7D, 0x52)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.06

sec = doc.sections[0]
sec.top_margin = Inches(0.6); sec.bottom_margin = Inches(0.6)
sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)

def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)

def ctext(cell, text, bold=False, color=None, size=8.5, align="left", white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(size)
    if white: r.font.color.rgb = WHITE
    elif color is not None: r.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "8"); b.set(qn("w:space"), "3"); b.set(qn("w:color"), "0B2545")
    pbdr.append(b); pPr.append(pbdr)
    return p

def h2(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BLUE
    return p

def para(text, size=10, bold=False, italic=False, color=None, align="just", space_after=6):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "just": WD_ALIGN_PARAGRAPH.JUSTIFY,
                   "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    return p

def bullet(runs):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    for t, b in runs:
        r = p.add_run(t); r.bold = b; r.font.size = Pt(10)
    return p

def table(headers, rows, widths=None, header_bg="0B2545", first_bold=True, font=8.3, hdr_align_first="left"):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    hc = t.rows[0].cells
    for i, hh in enumerate(headers):
        shade(hc[i], header_bg)
        ctext(hc[i], hh, bold=True, white=True, size=font, align=hdr_align_first if i == 0 else "center")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        if ri % 2 == 1:
            for c in cells: shade(c, "EEF2F7")
        for i, v in enumerate(row):
            ctext(cells[i], v, bold=(first_bold and i == 0), size=font, align="left" if i == 0 else "center")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def chart(fname, caption, width=6.6, src="Source: Company, United Spirits Excell.xlsx, Analyst estimates"):
    path = os.path.join(CH, fname)
    if not os.path.exists(path): return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after = Pt(7)
    r = c.add_run(caption + "   |   " + src); r.italic = True; r.font.size = Pt(7.5); r.font.color.rgb = GREY

# =============================================================== PAGE 1: COVER
for _ in range(3): doc.add_paragraph()
para("EQUITY RESEARCH  |  STOCK NOTE", size=12, bold=True, color=GOLD, align="center", space_after=2)
para("Sector: Alcoholic Beverages (Alcobev)  |  India", size=10, color=GREY, align="center", space_after=18)
para("UNITED SPIRITS LIMITED", size=30, bold=True, color=NAVY, align="center", space_after=2)
para("A Diageo Group Company", size=13, italic=True, color=GREY, align="center", space_after=24)
banner = doc.add_table(rows=1, cols=3); banner.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (lbl, val, clr) in enumerate([("RECOMMENDATION", "HOLD", NAVY),
                                     ("BASE CASE FAIR VALUE", "Rs 1,348", GREEN),
                                     ("BULL CASE FAIR VALUE", "Rs 1,500", GREEN)]):
    cell = banner.rows[0].cells[i]; shade(cell, "0B2545" if i == 0 else "EEF2F7")
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = cell.paragraphs[0].add_run(lbl + "\n"); r1.font.size = Pt(8); r1.bold = True
    r1.font.color.rgb = WHITE if i == 0 else GREY
    r2 = cell.paragraphs[0].add_run(val); r2.font.size = Pt(15); r2.bold = True
    r2.font.color.rgb = WHITE if i == 0 else clr
    cell.width = Inches(2.3)
doc.add_paragraph().paragraph_format.space_after = Pt(10)
para("Current Market Price: Rs 1,343        Add on dips band: Rs 1,150 to Rs 1,200        Time Horizon: 12 months",
     size=10.5, bold=True, color=NAVY, align="center", space_after=20)
para("Analyst: Equity Research (CFA)", size=10, bold=True, align="center", space_after=1)
para("Report Date: 26 June 2026", size=10, color=GREY, align="center", space_after=16)
dt = doc.add_table(rows=1, cols=1); dt.style = "Table Grid"; dt.alignment = WD_TABLE_ALIGNMENT.CENTER
shade(dt.rows[0].cells[0], "FBF6E9")
dp = dt.rows[0].cells[0].paragraphs[0]
rr = dp.add_run("For academic and illustrative use only. "); rr.bold = True; rr.font.size = Pt(8); rr.font.color.rgb = GOLD
rr2 = dp.add_run("This Stock Note is prepared for an MBA Finance capstone and follows CFA Institute research reporting essentials. "
                 "It is not investment advice and not an offer to buy or sell any security. Financial data is taken from the uploaded model "
                 "(United Spirits Excell.xlsx), the company filing (fy 2026.xlsx) and the FY2025 Annual Report. Industry, quarterly and macro "
                 "information is taken from public sources that are cited in the text. Third party content has been put in our own words to respect licensing rules.")
rr2.font.size = Pt(8); rr2.font.color.rgb = GREY
doc.add_page_break()

# =============================================================== PAGE 2: SNAPSHOT + OUR TAKE
h1("United Spirits Ltd  |  Stock Note")
# top recommendation strip
table(["Industry", "CMP (Rs)", "Recommendation", "Base Case FV", "Bull Case FV", "Horizon"],
      [["Alcobev", "1,343", "HOLD; add on dips in Rs 1,150-1,200 band", "Rs 1,348", "Rs 1,500", "12 months"]],
      widths=[0.9, 0.8, 2.4, 0.95, 0.95, 0.9], font=8.2)

# two columns feel: stock data table then Our Take
h2("Stock Data")
table(["Particulars", "Value", "Particulars", "Value"],
      [["NSE Code", "UNITDSPR", "Market Cap (Rs cr)", "96,453"],
       ["BSE Code", "532432", "Enterprise Value (Rs cr)", "93,726"],
       ["Bloomberg", "UNSP IN", "52 Week High (Rs)", "1,489"],
       ["Face Value (Rs)", "2", "52 Week Low (Rs)", "1,210"],
       ["Equity Shares O/S (cr)", "72.74", "Promoter (Diageo) holding", "About 56%"],
       ["Equity Capital (Rs cr)", "145", "Public holding", "About 44%"]],
      widths=[1.9, 1.45, 1.9, 1.45], font=8.3)

h2("Our Take")
para("United Spirits Limited (USL) is the largest beverage alcohol company in India. It makes, sells and distributes spirits such as whisky, "
     "brandy, rum, vodka and gin through a large set of Indian and global brands. USL is the India operating company of Diageo plc, the world "
     "leader in premium spirits, which owns about 56% of the equity. The company sells more than 60 brands, several of which sell over a million "
     "cases a year, and reaches consumers through more than 70,000 retail outlets and about 40 manufacturing sites across India.", align="just")
para("USL divides its brands into two groups. Prestige and Above (P&A) is the premium part of the portfolio and now makes up close to 90% of net "
     "sales. Popular is the mass market part, which the company has reduced on purpose to focus on premium growth. The company has spent the last "
     "few years reshaping the portfolio towards luxury, premium and upper prestige brands. This is the heart of the investment story, because in "
     "India premium spirits are growing faster than mass market spirits, and premium brands earn much higher margins.", align="just")
para("The India spirits market has strong long term support. Per person alcohol consumption is still low, the population is young, cities are "
     "growing, incomes are rising and people are choosing better quality drinks. United Spirits is well placed to benefit because it is the clear "
     "leader in Indian whisky and has the widest premium portfolio. Two recent events stand out. In March 2026 USL sold its entire stake in the "
     "Royal Challengers Bengaluru (RCB) cricket franchise for Rs 16,660 crore in an all cash deal, which unlocks a large amount of value. The "
     "company also completed the purchase of craft gin maker NAO Spirits, which strengthens its presence in fast growing new categories.",
     align="just")
para("Our view: USL is a high quality, market leading business with a strong balance sheet and rising margins. The shares, however, already trade "
     "at rich multiples (about 57 times FY27 estimated earnings). The base case fair value is close to the current price, so the risk and reward "
     "look balanced. We rate the stock HOLD and would add on declines into the Rs 1,150 to Rs 1,200 band.", align="just", bold=False)
doc.add_page_break()

# =============================================================== PAGE 3: VALUATION & RECOMMENDATION + FIN SUMMARY
h1("Valuation and Recommendation")
para("Over the last few years United Spirits has reshaped its portfolio towards premium brands. The share of Prestige and Above in net sales has "
     "risen from about 66% several years ago to close to 90% today. The company keeps adding new premium products such as the Godawan single malt, "
     "Royal Challenge American Pride and flavoured spirits, and it has renovated older brands like Antiquity and Black Dog. We expect premiumisation "
     "and new launches to keep driving value led revenue growth of about 9% to 10% a year over the next two to three years, with margins improving "
     "slowly as the mix gets richer and as cost saving programmes help.", align="just")
para("We value USL using a blend of methods, which is the approach used in the uploaded model. We give 60% weight to EV to EBITDA, 20% to price to "
     "earnings, 5% to EV to EBIT and 15% to a discounted cash flow (DCF). The relative methods carry more weight because there are several listed "
     "Indian alcobev peers whose market multiples are useful and current, while the long horizon DCF is more sensitive to assumptions. This blend gives "
     "a base case fair value of Rs 1,348 per share, which is about 57 times FY27 estimated earnings. Our bull case fair value is Rs 1,500, about 64 "
     "times FY27 estimated earnings, which would apply if premiumisation accelerates and margins rise faster. Because the base case is close to the "
     "current price of Rs 1,343, we rate the stock HOLD and suggest adding on declines in the Rs 1,150 to Rs 1,200 band.", align="just")
para("Important point on the RCB sale: our fair value is built from the core spirits business and does not add the Rs 16,660 crore of cash from the "
     "RCB sale. Depending on how this cash is used, for example a special dividend, a buyback or reinvestment, it is a source of extra value above "
     "our base case. This is one reason we prefer HOLD rather than a more negative call.", align="just")

h2("Financial Summary")
table(["Particulars (Rs cr)", "Q4FY26", "Q4FY25", "YoY %", "Q3FY26", "QoQ %", "FY25", "FY26", "FY27E", "FY28E"],
      [["Net Sales Value", "3,046", "2,946", "3.4", "3,683", "-17.3", "12,069", "12,467", "13,589", "14,812"],
       ["EBITDA", "591", "505", "17.0", "618", "-4.4", "2,236", "2,279", "2,547", "2,851"],
       ["APAT", "539", "421", "28.0", "418", "28.9", "1,582", "1,838", "1,709", "1,891"],
       ["EPS (Rs)", "7.4", "5.8", "", "5.7", "", "21.8", "25.3", "23.5", "26.0"],
       ["EBITDA margin %", "19.4", "17.1", "", "16.8", "", "18.5", "18.3", "18.7", "19.2"],
       ["P/E (x)", "", "", "", "", "", "62", "53", "57", "52"],
       ["EV/EBITDA (x)", "", "", "", "", "", "42", "41", "33", "30"]],
      widths=[1.55, 0.72, 0.72, 0.6, 0.72, 0.62, 0.72, 0.72, 0.72, 0.72], font=7.8)
para("Note: Quarterly NSV and EBITDA are on a standalone basis as reported by the company; APAT is the consolidated number widely reported. Annual "
     "figures are consolidated and follow the uploaded model. FY27 estimated EPS looks lower than FY26 because FY26 profit included one off items; "
     "the underlying business keeps growing. P/E and EV/EBITDA are on the current market price of Rs 1,343. Source: Company filings, United Spirits "
     "Excell.xlsx.", size=8, italic=True, color=GREY)
doc.add_page_break()

# =============================================================== PAGE 4: CHARTS IN FOCUS
h1("Charts in Focus")
chart("ex1_nsv_growth.png", "Net Sales Value and growth", width=6.5)
chart("ex3_volume_mix.png", "Volume mix: P&A versus Popular", width=6.0)
chart("ex2_margins.png", "Gross, EBITDA, EBIT and PAT margins", width=6.5)
doc.add_page_break()
chart("ex4_returns.png", "Return on equity and return on capital employed", width=6.5)
para("What the charts show: Net sales keep rising even though total volume has fallen, because the company sells more premium product at higher "
     "prices. The Prestige and Above share of volume has climbed steadily while Popular has shrunk. Gross margin has recovered to about 46% to 47% "
     "and EBITDA margin has moved up towards 20%. Returns on equity and on capital employed have stayed strong at about 20% or more, which shows "
     "the business uses capital efficiently.", align="just")
doc.add_page_break()

# =============================================================== PAGE 5: INDUSTRY ANALYSIS
h1("Industry Analysis: Indian Alcobev")
para("India is the third largest alcoholic beverage market in the world by volume and the largest market for whisky. Total Indian Made Foreign "
     "Liquor (IMFL) volume is about 410 million cases a year (IWSR, 2024). The wider alcohol market is worth roughly 50 to 70 billion US dollars and "
     "is growing at about 6% a year. Indian made spirits, both IMFL and country liquor, make up close to 70% of all alcohol consumed; beer is the "
     "other large category. The premium end is led by two global groups, Diageo (through United Spirits) and Pernod Ricard. Sources: IWSR via "
     "Livemint; industry reports.", align="just")
chart("dgm_industry_size.png", "Market size and consumption mix", width=6.8, src="Source: Industry reports, IWSR via public media")
para("The chart on the left shows the market growing from about 52 billion US dollars in 2021 towards an estimated 77 billion US dollars by 2030. "
     "The pie on the right shows that Indian spirits dominate consumption. This is helpful for USL because it leads in IMFL and has a pan India network.",
     align="just")
chart("dgm_value_chain.png", "How the industry works and where the profit sits", width=7.0, src="Source: Analyst, based on industry structure")
para("The value chain diagram explains the flow from raw materials to the consumer. The single biggest cost in the chain is state excise duty, which "
     "is collected from consumers and passed to state governments. United Spirits captures the most value in premium blending and in the strength of "
     "its brands. Because each state controls its own taxes and rules, the topline is best tracked as Net Sales Value, which is revenue after excise.",
     align="just")
doc.add_page_break()

# =============================================================== PAGE 6: PREMIUMISATION
h1("Premiumisation: The Main Growth Driver")
para("Premiumisation means consumers trading up to better and more expensive drinks. It is the most important trend for United Spirits. Research "
     "shows that premium IMFL is growing at low double digit rates while mass market brands grow at mid single digit rates (JM Financial, September "
     "2025, via Financial Express). Luxury spirits are only about 3% of cases sold but contribute more than 20% of industry profits (IWSR via "
     "Livemint). This matters a lot for USL because close to 90% of its net sales now come from Prestige and Above brands, so every step up in the "
     "mix adds strongly to gross and EBITDA margins.", align="just")
chart("dgm_premium_pyramid.png", "The premiumisation pyramid: price ladder and growth", width=6.6, src="Source: Analyst; growth rates per company and industry data")
para("The pyramid shows the price ladder. Luxury sits at the top with the highest growth, followed by premium and prestige, while popular sits at "
     "the bottom with flat to falling growth. United Spirits has moved most of its sales up the pyramid. The financial result is higher realisations "
     "per case and steadily rising margins, even when total volume growth is modest.", align="just")
para("Why this is good for profits: a richer mix lifts gross margin, and because many costs do not rise as fast as premium revenue, more of each "
     "extra rupee of sales drops to EBITDA. This is the engine behind our forecast of EBITDA margin moving from about 18% towards 20% over the next "
     "few years.", align="just")
doc.add_page_break()

# =============================================================== PAGE 7: PORTER + REGULATION
h1("Competitive Forces and Regulation")
chart("dgm_porter.png", "Porter's Five Forces for Indian spirits", width=6.6, src="Source: Analyst framework")
para("Threat of new entrants is low. A new national player would need to deal with more than 30 separate state licensing and tax systems, build "
     "in state supply chains, invest in whisky ageing and overcome strict rules that ban direct alcohol advertising. Supplier power is moderate; key "
     "inputs such as extra neutral alcohol, scotch concentrate and glass can rise in price at times. Buyer power is high because in many states the "
     "government corporations control pricing and distribution, and one large state customer is close to 30% of revenue. The threat of substitutes "
     "(beer, wine and low or no alcohol drinks) is moderate. Rivalry is high and rising, with Pernod Ricard, Radico Khaitan, Allied Blenders and a "
     "newly strengthened Tilaknagar all pushing premium brands.", align="just")
h2("State Regulation and Taxes")
para("Alcohol is outside GST and is controlled by each state. States set licences, price approvals, route to market and excise duty. The most "
     "important recent change is in Karnataka. From 11 May 2026 Karnataka became the first state to move to an Alcohol in Beverage based excise "
     "system, which taxes drinks by their actual alcohol content and frees producers to set prices within set bands rather than under fixed government "
     "prices (Economic Times; Hindustan Times). Karnataka is the largest whisky market in India at about 17% of national volumes, so more pricing "
     "freedom there is a structural positive for premium players like USL, although the change brings some short term uncertainty. The main downside "
     "risks remain dry states and sudden changes to route to market rules.", align="just")
doc.add_page_break()

# =============================================================== PAGE 8: MACRO
h1("Economy and Demand Drivers")
table(["Indicator", "Reading", "Source"],
      [["India real GDP growth FY26", "7.7% (Q4 7.8%); nominal 8.9%", "MoSPI / Economic Survey"],
       ["FY27 GDP estimate", "6.8% to 7.2% (Survey); RBI about 7.4%", "PRS / RBI"],
       ["Position", "Fastest growing major economy, fourth year", "Economic Survey 2025-26"],
       ["Demographics", "Young legal drinking age base; fast urbanisation", "Census / industry"]],
      widths=[2.4, 3.1, 1.5], font=8.4)
para("How the economy feeds spirits demand: India is growing fast, the middle class is expanding, incomes are rising, cities are growing and the "
     "population is young. These factors support steady growth in branded premium spirits. Importantly, the trend is towards drinking better rather "
     "than drinking more, which suits USL because its sales are concentrated in premium brands. The main risk is a slowdown in discretionary "
     "spending, for example from high food inflation or weak urban sentiment, which can cause down trading where some consumers move to cheaper "
     "products. Sources: Livemint; Economic Times; Economic Survey 2025-26.", align="just")

h1("Competitive Positioning")
chart("dgm_positioning.png", "Where the players sit: premium mix versus margin", width=6.4, src="Source: Company FY26 data; bubble size is EBITDA")
para("The map places each company by how premium its portfolio is (left to right) and by its FY26 EBITDA margin (bottom to top). United Spirits sits "
     "in the top right, with the most premium portfolio and a high margin. Radico Khaitan is close behind on premium positioning. Allied Blenders and "
     "Tilaknagar are more mid market, and United Breweries is shown as a beer reference with a lower margin. The bubble size reflects EBITDA in rupees.",
     align="just")
doc.add_page_break()

# =============================================================== PAGE 9: QUARTERLY OVERVIEW
h1("Quarterly Performance and Concall Analysis: FY26")
para("This section reviews the four quarters of FY26 (Q1 ended June 2025 to Q4 ended March 2026) using company results and management commentary "
     "from the earnings calls. The uploaded Excel model does not contain quarterly numbers, so the quarterly data here is taken from company filings "
     "and public reports (Diageo India results releases; Economic Times; CNBC TV18; Business Standard; Quartr). Net sales value (NSV) and EBITDA are "
     "on a standalone basis, and profit after tax (PAT) is the consolidated figure. The four quarters together add up to the full year standalone NSV "
     "of Rs 12,448 crore and standalone EBITDA of Rs 2,296 crore.", align="just")
table(["Rs cr", "Q1FY26", "Q2FY26", "Q3FY26", "Q4FY26", "FY26"],
      [["Net Sales Value", "2,549", "3,170", "3,683", "3,046", "12,448"],
       ["NSV growth YoY %", "8.4", "11.5", "7.6", "3.4", "7.6"],
       ["EBITDA", "415", "672", "618", "591", "2,296"],
       ["EBITDA margin %", "16.3", "21.2", "16.8", "19.4", "18.4"],
       ["PAT (consolidated)", "417", "464", "418", "539", "1,838"],
       ["PAT growth YoY %", "-14", "36", "25", "28", "Mixed"]],
      widths=[1.7, 0.95, 0.95, 0.95, 0.95, 0.95], font=8.2)
chart("q_nsv_margin.png", "Quarterly NSV and EBITDA margin", width=6.2, src="Source: Company results; Diageo India releases")
chart("q_growth.png", "Quarterly growth: NSV versus PAT", width=6.2, src="Source: Company results")
doc.add_page_break()

chart("q_ebitda_pat.png", "Quarterly EBITDA and PAT", width=6.2, src="Source: Company results")
h2("Quarterly story in plain words")
para("FY26 was a year of two halves in terms of margins. The first quarter was weak on profit, the second quarter was very strong, the third quarter "
     "saw soft topline growth during the festive season, and the fourth quarter delivered a strong margin and profit even though sales growth was "
     "slow. Across the year the premium portfolio kept growing and the company held its strategy of value over volume. The big strategic news came in "
     "the fourth quarter and just after: the completion of the NAO Spirits purchase and the approval and sale of the RCB cricket franchise. Management, "
     "led by CEO and Managing Director Praveen Someshwar, repeated that the focus is double digit medium term revenue growth with margins slightly "
     "ahead.", align="just")
doc.add_page_break()

# =============================================================== PAGE 10-11: PER QUARTER DETAIL
h1("Earnings Call Highlights by Quarter")

h2("Q1 FY26 (quarter ended June 2025): profit dipped on mix and costs")
bullet([("Result: ", True), ("Standalone NSV grew about 8.4% to Rs 2,549 crore, helped by the re-entry into the Andhra Pradesh market. EBITDA fell "
        "about 9% to Rs 415 crore and the margin dropped to 16.3% from 19.5% a year earlier. Consolidated PAT fell about 14% to Rs 417 crore "
        "(CNBC TV18; Economic Times).", False)])
bullet([("Why margins fell: ", True), ("the Andhra Pradesh re-entry added a large amount of lower margin Popular volume, input costs were higher, and "
        "the company reinvested about 9.3% of NSV in advertising and promotion to support key brands. So sales rose but profit fell.", False)])
bullet([("Management view: ", True), ("the premium segment kept growing and the company framed the margin dip as a timing and mix effect rather than a "
        "change in the long term path. It maintained its guidance of double digit medium term growth.", False)])

h2("Q2 FY26 (quarter ended September 2025): a strong beat across the board")
bullet([("Result: ", True), ("standalone NSV rose about 11.5% to Rs 3,170 crore. EBITDA jumped about 32.5% to Rs 672 crore and the margin reached "
        "21.2%. Consolidated PAT rose about 36% to Rs 464 crore. The shares rose about 7% on the day (Diageo India release; CNBC TV18; Economic Times).",
        False)])
bullet([("Drivers: ", True), ("Prestige and Above NSV grew about 12.4% and made up 89.6% of net sales. Even the Popular segment grew about 9.2%. "
        "Premium brands such as Johnnie Walker, Signature and Antiquity led the growth, and operating leverage lifted the margin sharply (Storyboard18).",
        False)])
bullet([("Management view: ", True), ("management pointed to broad based premiumisation and good cost control. This quarter showed the full power of the "
        "premium mix when costs are well managed.", False)])

h2("Q3 FY26 (quarter ended December 2025): soft topline in the festive quarter")
bullet([("Result: ", True), ("this is seasonally the biggest quarter. Standalone NSV grew about 7.6% to Rs 3,683 crore. Reported EBITDA was about Rs 599 "
        "to 618 crore, with growth of around 5.5% as the company stepped up advertising during the festive and wedding season. Consolidated PAT rose "
        "about 25% to Rs 418 crore. The board declared an interim dividend of Rs 6 per share. The shares fell about 3% as investors had hoped for "
        "faster sales growth (Economic Times; Businessworld).", False)])
bullet([("Management view: ", True), ("management said about 80% of the national portfolio grew at double digit rates. It announced a relaunch of "
        "McDowell's No.1 to win back share in the lower Prestige segment, showing a focus on the parts of the portfolio that had lagged (Quartr).",
        False)])

h2("Q4 FY26 (quarter ended March 2026): strong margin, slow sales, big strategic moves")
bullet([("Result: ", True), ("standalone NSV grew about 3.4% to Rs 3,046 crore, slowed by a high base and a seasonally weak quarter. Prestige and Above "
        "NSV grew about 5% to Rs 2,745 crore. EBITDA rose about 17% to Rs 591 crore and gross margin reached 47.3%, a multi quarter high, helped by "
        "lower advertising and employee costs. Consolidated PAT rose about 28% to Rs 539 crore. A final dividend was declared (Diageo India release; "
        "NDTV Profit; Businessworld, citing JPMorgan).", False)])
bullet([("Strategic moves: ", True), ("the company completed the acquisition of NAO Spirits (the maker of Greater Than and Hapusa gin) and approved the "
        "sale of the RCB cricket franchise. The RCB sale for Rs 16,660 crore was completed in March 2026 to a group led by the Aditya Birla Group "
        "(Quartr; The Hindu).", False)])
bullet([("Management view: ", True), ("management struck a positive tone on the growth outlook for FY27 and reiterated its focus on premiumisation, "
        "innovation and productivity.", False)])
doc.add_page_break()

# =============================================================== PAGE 12: INVESTMENT RATIONALE / PORTFOLIO
h1("Investment Rationale")
h2("A rich and broad brand portfolio is the strongest pillar of growth")
para("United Spirits owns and operates one of the widest spirits portfolios in India, with more than 60 brands across scotch, IMFL whisky, brandy, "
     "rum, vodka and gin. Several brands sell over a million cases a year. The portfolio mixes old and new, large and small, and global and local "
     "brands. The company also imports and sells many famous Diageo brands such as Johnnie Walker, Captain Morgan, Baileys, Singleton, Tanqueray, "
     "Don Julio, Ketel One and Smirnoff under licence. This spread across price points lets USL serve both aspiring and affluent consumers and gives "
     "it pricing power at the premium end.", align="just")
table(["Tier", "Marquee Brands", "Indicative Price (Rs/case)", "Indicative 5-yr CAGR"],
      [["Luxury", "Johnnie Walker, Godawan, Singleton, Tanqueray, Don Julio", "Above 2,000", "About 30% to 35%"],
       ["Premium", "Black and White, Black Dog, Signature, Smirnoff", "800 to 2,000", "About 15%"],
       ["Prestige", "McDowell's No.1, Royal Challenge, Antiquity", "400 to 800", "About 5%"],
       ["Popular", "Director's Special Black, mass whiskies and rums", "Below 400", "Flat to negative"]],
      widths=[1.2, 3.0, 1.5, 1.4], font=8.2)
h2("A well timed pivot to premium supports double digit revenue growth")
para("In 2022 USL sold 32 Popular brands and franchised several more, so that it could focus on Prestige and Above. Since then the premium share of "
     "sales has climbed towards 90%. New launches such as Godawan single malt, which has won many international awards, and Royal Challenge American "
     "Pride, one of the fastest growing new products, have supported this shift. The renovation of brands such as Antiquity, Black Dog and McDowell's "
     "No.1 has also helped. This portfolio work is the main reason we expect value led growth to continue.", align="just")
h2("Strong cash and a clean balance sheet")
para("USL is effectively debt free and holds net cash. This gives it the flexibility to invest in brands, pay healthy dividends and absorb cost "
     "shocks. The Rs 16,660 crore from the RCB sale adds further to this strength and creates options for shareholder returns.", align="just")
doc.add_page_break()

# =============================================================== PAGE 13: FINANCIAL ANALYSIS
h1("Financial Analysis")
h2("Income Statement (consolidated, Rs cr)")
table(["Particulars", "FY24", "FY25", "FY26", "FY27E", "FY28E"],
      [["Net Sales Value", "11,321", "12,069", "12,467", "13,589", "14,812"],
       ["Growth %", "6.7", "6.6", "3.3", "9.0", "9.0"],
       ["Gross Profit", "4,779", "5,209", "5,793", "6,183", "6,814"],
       ["Gross Margin %", "42", "43", "46", "46", "46"],
       ["EBITDA", "2,000", "2,236", "2,279", "2,547", "2,851"],
       ["EBITDA Margin %", "18", "19", "18", "19", "19"],
       ["EBIT", "1,725", "1,953", "1,990", "2,170", "2,433"],
       ["PAT", "1,408", "1,582", "1,838", "1,709", "1,891"],
       ["PAT Margin %", "12", "13", "15", "13", "13"],
       ["EPS (Rs)", "19.4", "21.8", "25.3", "23.5", "26.0"]],
      widths=[2.1, 0.9, 0.9, 0.9, 0.9, 0.9], font=8.2)
para("The story in the income statement is clear: net sales rise every year even though total volume has fallen, because the mix keeps getting more "
     "premium. Gross margin recovered from a weak FY23 to about 46% by FY26 as input costs eased and the mix improved. EBITDA margin is moving up "
     "towards 20%. FY26 PAT was high partly due to one off items, which is why FY27 estimated PAT looks lower even though the underlying business keeps "
     "growing.", align="just")
h2("Balance Sheet (consolidated, Rs cr)")
table(["Particulars", "FY24", "FY25", "FY26"],
      [["Net Worth (Shareholders' funds)", "About 7,300", "About 8,200", "About 9,200"],
       ["Total Debt", "Low / negligible", "Low / negligible", "Low / negligible"],
       ["Cash and Treasury", "1,868", "2,903", "3,134"],
       ["Net Debt to EBITDA (x)", "-0.8", "-1.1", "-1.2"],
       ["Current Ratio (x)", "1.9", "2.0", "2.1"],
       ["Inventory days", "120", "116", "136"],
       ["Receivable days", "89", "98", "103"]],
      widths=[2.6, 1.3, 1.3, 1.3], font=8.3)
para("The balance sheet is a fortress. The company carries net cash, so net debt to EBITDA is negative. Working capital is heavy because premium "
     "whisky needs to age (high inventory) and because sales to state corporations stretch receivables. This is normal for the industry and is the main "
     "reason free cash conversion is moderate in any single year.", align="just")
doc.add_page_break()

# =============================================================== PAGE 14: CASH FLOW + RATIOS
h1("Cash Flow and Key Ratios")
h2("Cash Flow Summary (Rs cr)")
table(["Particulars", "FY26", "FY27E", "FY28E", "FY29E", "FY30E"],
      [["Operating Cash Flow", "1,459", "1,292", "2,345", "1,957", "2,831"],
       ["Capex", "(181)", "(204)", "(222)", "(240)", "(257)"],
       ["Free Cash Flow", "1,278", "1,088", "2,123", "1,717", "2,574"],
       ["Dividends Paid", "(1,263)", "(1,145)", "(1,267)", "(1,420)", "(1,553)"]],
      widths=[2.0, 1.0, 1.0, 1.0, 1.0], font=8.3)
chart("ex9_fcf_div.png", "Free cash flow versus dividends", width=5.8)
h2("Key Ratios")
table(["Ratio", "FY24", "FY25", "FY26", "FY27E", "FY28E"],
      [["ROE %", "21", "21", "22", "19", "19"],
       ["ROCE %", "23", "23", "21", "22", "22"],
       ["Gross Margin %", "42", "43", "46", "46", "46"],
       ["EBITDA Margin %", "18", "19", "18", "19", "19"],
       ["Net Debt to EBITDA (x)", "-0.8", "-1.1", "-1.2", "-1.1", "-1.0"],
       ["Dividend Payout %", "About 70", "About 70", "About 70", "About 70", "About 70"]],
      widths=[2.0, 0.9, 0.9, 0.9, 0.9, 0.9], font=8.3)
para("Capex is small, at about 1% to 2% of net sales, because the model is light on fixed assets. With healthy operating cash flow, low capex and a "
     "payout of about 70% of profit, USL is a steady cash compounder. The dip in ROE in FY27 reflects the one off items in the FY26 base.", align="just")
doc.add_page_break()

# =============================================================== PAGE 15-16: VALUATION DETAIL
h1("Valuation in Detail")
h2("A) Discounted Cash Flow (DCF)")
table(["Assumption", "Value"],
      [["Risk free rate (10 year G-Sec)", "7.0%"],
       ["Equity risk premium", "4.0%"],
       ["Adjusted beta", "0.92"],
       ["Cost of equity, which is also WACC (debt is negligible)", "11.0%"],
       ["Terminal growth", "6.0%"],
       ["Enterprise Value", "Rs 43,990 cr"],
       ["Add cash, less debt", "+3,134 / -407"],
       ["Equity Value", "Rs 46,717 cr"],
       ["DCF Value per share", "Rs 642"]],
      widths=[4.2, 2.4], font=8.4)
para("The DCF uses a WACC of 11% and a terminal growth of 6%, which is close to long run nominal consumption growth. On these inputs the DCF gives a "
     "value of about Rs 642 per share, which is below the market price. Heavy working capital lowers free cash conversion, and the market is clearly "
     "pricing in a lower discount rate or faster long term growth than our base case. For this reason the DCF acts as a floor and carries only a 15% "
     "weight in the blend.", align="just")
chart("ex7_dcf_sensitivity.png", "DCF value per share under different WACC and growth", width=5.0)

h2("B) Relative Valuation")
table(["Method", "Peer median (x)", "Implied value per share (Rs)"],
      [["EV/EBITDA", "39.2", "1,411"],
       ["P/E", "71.3", "1,675"],
       ["EV/EBIT", "45.9", "1,408"],
       ["EV/Revenue (reference)", "4.4", "859"]],
      widths=[2.3, 2.0, 2.3], font=8.4)
para("We apply the median forward multiples of the listed peers (Radico Khaitan, United Breweries and Allied Blenders) to USL's FY27 estimated "
     "figures. USL trades below the fastest growing peer, Radico, on P/E, but above mass market peers. This is fair given USL's leadership, premium "
     "mix and strong returns, while its slower headline growth versus Radico limits the premium.", align="just")
doc.add_page_break()

h2("C) Target Price: Blended Method")
table(["Method", "Weight", "Implied value (Rs)", "Contribution (Rs)"],
      [["EV/EBITDA", "60%", "1,411", "846"],
       ["P/E", "20%", "1,675", "335"],
       ["EV/EBIT", "5%", "1,408", "70"],
       ["DCF", "15%", "642", "96"],
       ["Base Case Fair Value", "100%", "", "1,348"]],
      widths=[1.8, 1.0, 1.7, 1.7], font=8.4)
chart("ex6_football_field.png", "Valuation summary by method and the blended target", width=6.4)
para("We give relative methods more weight because India has several listed alcobev peers whose multiples are useful and current, and because the long "
     "horizon DCF is very sensitive to small changes in assumptions. The blend gives a base case fair value of Rs 1,348. Our bull case of Rs 1,500 "
     "assumes faster premiumisation and margin gains, while our bear case of about Rs 920 assumes weak demand, margin pressure and a de-rating.",
     align="just")
chart("ex8_scenarios.png", "Bull, base and bear case fair values", width=4.8)
doc.add_page_break()

# =============================================================== PAGE 17: ESG
h1("ESG (Environment, Social and Governance)")
para("United Spirits runs its ESG agenda under the name Spirit of Progress, which is aligned with Diageo's global plan. The FY2025 Annual Report "
     "shows strong progress.", align="just")
h2("Environment")
bullet([("Water: ", True), ("about 54% improvement in water used per litre of spirit distilled since FY2020; programmes to replenish water and support "
        "regenerative farming.", False)])
bullet([("Energy and emissions: ", True), ("99% of energy in direct operations now comes from renewable sources; about 93% cut in direct Scope 1 and 2 "
        "greenhouse gas emissions since FY2020.", False)])
bullet([("Net Zero plan: ", True), ("Net Zero in direct operations by 2030 and across the value chain by 2050.", False)])
h2("Social")
bullet([("Responsible drinking: ", True), ("about 0.7 million people reached through responsible drinking programmes.", False)])
bullet([("People: ", True), ("about 2,400 employees, around 86% retention, about 89% engagement and about 28% women representation.", False)])
h2("Governance")
bullet([("", False), ("Diageo level governance standards, a strong board, a marketing code and a conservative treasury policy.", False)])
para("Impact on value and risk: strong ESG action lowers the risk of regulatory and reputation problems in a sector that is always under scrutiny, "
     "supports the premium brand image and is consistent with the low beta we use in the DCF. We do not add a separate ESG premium, but we see ESG "
     "leadership as supportive of a durable valuation. Source: United Spirits FY2025 Annual Report.", align="just")
doc.add_page_break()

# =============================================================== PAGE 18: RISKS
h1("Risks and Concerns")
table(["Risk", "Why it matters", "What reduces it"],
      [["Regulatory and route to market changes", "Sudden state policy or distribution changes can hurt sales and credit risk",
        "Spread across 30 plus states; Diageo compliance"],
       ["Excise and tax increases", "Higher duties raise shelf prices with little gain to the company",
        "Premium mix and pricing power; Karnataka pricing freedom"],
       ["Raw material inflation", "Extra neutral alcohol, glass and scotch can rise in price",
        "Scale buying, cost programmes, productivity"],
       ["Stronger competition", "Pernod, Radico, Allied Blenders and Tilaknagar are pushing premium",
        "Widest portfolio, deep distribution, value share focus"],
       ["Consumer down trading", "A slowdown can move some buyers to cheaper drinks",
        "Full price ladder; mass to premium funnel"],
       ["Currency and import cost", "Scotch concentrate is imported", "Mostly local production; some hedging"],
       ["Customer concentration", "One state customer is close to 30% of revenue", "Industry norm; careful receivables management"],
       ["Rich valuation", "High multiples leave little margin of safety", "Quality, growth and the RCB cash option"]],
      widths=[1.9, 2.7, 2.0], font=8.0)
h2("Global context")
para("USL is mainly a domestic business, so it has limited direct exposure to US tariffs. Even so, the global backdrop matters. Global spirits volumes "
     "fell about 2% in 2025, the third year of decline, and premiumisation has slowed in many markets (IWSR). US tariffs cut scotch exports to the US "
     "by about 15% by volume (Scotch Whisky Association via The Spirits Business). Parent Diageo cut its FY26 guidance on weak US and China demand and "
     "tariffs. The bright side is that India stands out as a growth market within a flat global industry, which is a relative positive for USL. The "
     "items to watch are scotch input costs and any change in the parent's strategy. Sources: IWSR; The Spirits Business; Diageo releases.", align="just")
doc.add_page_break()

# =============================================================== PAGE 19: COMPANY BACKGROUND + PEER TABLE
h1("Company Background and Peers")
para("United Spirits Ltd is the largest beverage alcohol company in India and the India operating company of Diageo plc, which holds about 56% of "
     "the equity. The company is based in Bengaluru. It makes, sells and distributes spirits across more than 60 brands, with about 40 manufacturing "
     "sites and more than 70,000 retail outlets. The company also imports and sells many global Diageo brands under licence. It reports two main "
     "segments: beverage alcohol (the core business) and, until March 2026, a sports segment (the RCB cricket franchise), which has now been sold. "
     "The Chief Executive Officer and Managing Director is Praveen Someshwar.", align="just")
h2("Peer Comparison (FY26)")
table(["Company", "NSV (Rs cr)", "EBITDA (Rs cr)", "EBITDA %", "Mcap (Rs cr)", "EV/EBITDA", "P/E"],
      [["United Spirits", "12,467", "2,279", "18.3", "96,453", "41", "53"],
       ["Radico Khaitan", "6,050", "1,038", "17.1", "50,274", "49", "82"],
       ["United Breweries (beer)", "9,240", "824", "8.9", "35,605", "44", "86"],
       ["Allied Blenders", "3,923", "542", "13.8", "17,640", "34", "77"]],
      widths=[1.8, 1.0, 1.05, 0.85, 1.05, 0.85, 0.7], font=8.0)
chart("ex5_peer_multiples.png", "Peer valuation multiples", width=6.0, src="Source: Company data, United Spirits Excell.xlsx")
para("USL has the largest premium portfolio, the deepest distribution and the strongest balance sheet in the group. It trades below Radico on P/E "
     "because Radico is growing faster from a smaller base, while it trades above mass market peers because of its leadership and mix.", align="just")
doc.add_page_break()

# =============================================================== PAGE 20: RATING + DISCLAIMER
h1("Recommendation Summary and Disclaimer")
table(["Item", "Detail"],
      [["Recommendation", "HOLD; add on dips in the Rs 1,150 to Rs 1,200 band"],
       ["Current Market Price", "Rs 1,343"],
       ["Base Case Fair Value", "Rs 1,348 (about 57 times FY27 estimated EPS)"],
       ["Bull Case Fair Value", "Rs 1,500 (about 64 times FY27 estimated EPS)"],
       ["Bear Case Fair Value", "About Rs 920"],
       ["Time Horizon", "12 months"],
       ["Key catalysts", "Use of RCB cash; Karnataka pricing; premium volume; input costs"]],
      widths=[2.0, 4.6], font=8.6)
para("In one paragraph: United Spirits is a high quality, market leading business with a premium portfolio, a fortress balance sheet, strong returns "
     "and a clear premiumisation runway. The shares already trade at rich multiples, so the base case fair value sits close to the current price. We "
     "rate the stock HOLD and would add on declines. The Rs 16,660 crore RCB cash, not included in our base case, is an extra source of value that "
     "supports our preference for HOLD over a more negative call. The stock deserves a premium to mass market peers for its quality and mix, but a "
     "discount to the fastest growing peer (Radico) given its larger, more mature base.", align="just")
h2("Rating scale")
para("HOLD means the expected return over the next 12 months is broadly in line with the broad market. BUY means meaningfully above the market, and "
     "SELL means meaningfully below the market. Add on dips means accumulate the stock if it falls into the stated price band.", size=9.5)
h2("Disclaimer")
para("This Stock Note has been prepared for an MBA Finance capstone and for illustration, following CFA Institute research reporting essentials. It is "
     "not investment advice and is not an offer or solicitation to buy or sell any security. Readers should take their own professional advice before "
     "investing. Financial data is taken from the uploaded model (United Spirits Excell.xlsx), the company filing (fy 2026.xlsx) and the FY2025 Annual "
     "Report. Industry, quarterly and macro information is taken from public sources that are cited in the text, including Diageo India results "
     "releases, IWSR, Economic Times, Livemint, CNBC TV18, Business Standard, Businessworld, NDTV Profit, The Hindu, Hindustan Times, The Spirits "
     "Business, Quartr, PRS and the Economic Survey. Third party content has been summarised in our own words to respect licensing rules. The author "
     "holds no position in the stock.", size=9, align="just")

# ---- footer with page numbers ----
def add_page_number(p):
    run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    ins = OxmlElement('w:instrText'); ins.set(qn('xml:space'), 'preserve'); ins.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(ins); run._r.append(f2)

footer = doc.sections[0].footer
fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("United Spirits Ltd  |  Stock Note  |  HOLD, Base Case FV Rs 1,348  |  Page ")
fr.font.size = Pt(7.5); fr.font.color.rgb = GREY
add_page_number(fp)

out = os.path.join(HERE, "United_Spirits_Stock_Note.docx")
doc.save(out)
print("SAVED:", out, os.path.getsize(out), "bytes")
